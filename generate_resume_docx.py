import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume_docx(input_md_path, output_docx_path):
    if not os.path.exists(input_md_path):
        print(f"Error: {input_md_path} not found.")
        return

    doc = Document()

    # Adjust Margins for a one-page layout
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    with open(input_md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        clean_line = line.strip()
        if not clean_line:
            continue

        # Header 1: Name
        if clean_line.startswith('# '):
            p = doc.add_heading(clean_line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(24)
            run.font.bold = True

        # Header 2: Section Titles
        elif clean_line.startswith('## '):
            p = doc.add_heading(clean_line[3:], level=2)
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
            # Add a horizontal line or spacing if needed
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        # Subtitles (Bold lines like Roles/Contact)
        elif clean_line.startswith('**') and clean_line.endswith('**'):
            text = clean_line.replace('**', '')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(12)

        # Bullet Points
        elif clean_line.startswith('- '):
            text = clean_line[2:]
            # Handle bold inside bullets (e.g., **TypeScript**)
            p = doc.add_paragraph(style='List Bullet')
            parts = text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.font.bold = True
            p.paragraph_format.space_after = Pt(2)

        # Experience Header Lines (Company | Role | Dates)
        elif '|' in clean_line and clean_line.count('**') >= 2:
            text = clean_line.replace('**', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(6)

        # General Text / Links
        else:
            # Remove markdown links [Text](URL) for a cleaner Word look
            import re
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_line)
            
            p = doc.add_paragraph()
            # Center contact info if it contains email/phone
            if '@' in text or '|' in text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            parts = text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.font.bold = True
            
            p.paragraph_format.space_after = Pt(2)

    doc.save(output_docx_path)
    print(f"Successfully created: {output_docx_path}")

if __name__ == "__main__":
    input_path = "/Users/ilazarkov/Development/VSCodeProjects/prototyping/my-github-profile/Resume-one-page.md"
    output_path = "/Users/ilazarkov/Development/VSCodeProjects/prototyping/my-github-profile/Ihor_Lazarkov_Resume.docx"
    create_resume_docx(input_path, output_path)