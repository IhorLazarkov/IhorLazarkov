import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

def convert_md_to_docx():
    # Set paths relative to this script's location
    script_dir = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(script_dir, "../input/CoverLetter.md")
    output_path = os.path.join(script_dir, "../output/Renew_Home_CL.docx")
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    if not os.path.exists(input_path):
        print(f"Error: Input file not found at {input_path}")
        return

    print(f"Generating Word document from {input_path} using python-docx...")
    
    doc = Document()
    
    # Document Styling: Set margins to 0.5 inches
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue

        # Handle Headings
        if stripped_line.startswith("# "):
            p = doc.add_heading(stripped_line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped_line.startswith("## "):
            doc.add_heading(stripped_line[3:], level=1)
        # Handle List Items
        elif stripped_line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, stripped_line[2:])
        # Handle Regular Paragraphs
        else:
            p = doc.add_paragraph()
            # Center the contact/intro line if it's near the top
            if "|" in stripped_line and len(doc.paragraphs) < 5:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_formatted_text(p, stripped_line)

    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

def add_hyperlink(paragraph, url, text):
    """
    Adds a hyperlink to a paragraph object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def _add_formatted_text(paragraph, text):
    """Processes basic Markdown bold and links."""
    # Split by bold tags (**...**) and link tags ([text](url))
    parts = re.split(r'(\*\*.*?\*\*|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('[') and '](' in part:
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                add_hyperlink(paragraph, match.group(2), match.group(1))
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    convert_md_to_docx()
